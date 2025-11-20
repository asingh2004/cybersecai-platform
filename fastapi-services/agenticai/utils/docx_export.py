import os
import uuid
from fastapi import HTTPException
from fastapi.responses import FileResponse

# Safe logger fallback (avoids NameError if utils.logging is unavailable)
try:
    from utils.logging import log_to_laravel as _log
except Exception:
    def _log(msg: str):
        try:
            print(f"[docx_export] {msg}")
        except Exception:
            pass

BASE_DIR = "/home/cybersecai/htdocs/www.cybersecai.io/webhook/agenticai/tmp_csv"


def download_docx(file: str):
    # Sanitize filename and ensure .docx
    if not file or not file.lower().endswith(".docx"):
        raise HTTPException(status_code=404, detail="File not found")

    filename = os.path.basename(file)
    if filename != file:
        # Prevent path traversal
        raise HTTPException(status_code=404, detail="File not found")

    file_path = os.path.join(BASE_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )


def markdown_to_docx(md_content: str, docx_path: str):
    import re
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    def add_paragraph_safe(document, text, style_name=None):
        try:
            if style_name:
                return document.add_paragraph(text, style=style_name)
            return document.add_paragraph(text)
        except Exception:
            # Fallback to default style if the style doesn't exist
            return document.add_paragraph(text)

    # Ensure we have a string
    if md_content is None:
        md_content = ""
    elif not isinstance(md_content, str):
        md_content = str(md_content)

    doc = Document()
    lines = md_content.splitlines()
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Handle fenced code blocks to avoid mis-parsing tables/bullets within
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            # Render fenced blocks as simple paragraphs for now
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                add_paragraph_safe(doc, lines[i])
                i += 1
            # Move past closing fence if present
            i += 1
            continue

        if in_code_block:
            add_paragraph_safe(doc, line)
            i += 1
            continue

        # Headings (# ... ###### ...)
        if re.match(r"^#{1,6} ", line):
            level = len(line.split(" ")[0])  # count of '#'
            text = line[level + 1:].strip() if len(line) > level + 1 else ""
            try:
                doc.add_heading(text, min(max(level - 1, 0), 4))
            except Exception:
                add_paragraph_safe(doc, text)

        # Tables starting with '|' (parse contiguous block)
        elif line.strip().startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Backtrack one step because outer loop will increment
            i -= 1

            try:
                header_line = table_lines[0].strip().strip("|")
                headers = [c.strip() for c in header_line.split("|")]
                if not headers:
                    # Fallback to plain paragraphs if header empty
                    for tl in table_lines:
                        add_paragraph_safe(doc, tl)
                else:
                    table = doc.add_table(rows=1, cols=len(headers))
                    try:
                        table.style = "Light Grid Accent 1"
                    except Exception:
                        pass
                    for idx, text in enumerate(headers):
                        try:
                            table.rows[0].cells[idx].text = text
                        except Exception:
                            pass
                    # Data rows: skip the separator row if present (usually second line)
                    data_start_idx = 1
                    if len(table_lines) > 1 and set(table_lines[1].replace("|", "").strip()) <= set("-: "):
                        data_start_idx = 2
                    for trow in table_lines[data_start_idx:]:
                        vals = [c.strip() for c in trow.strip().strip("|").split("|")]
                        row = table.add_row()
                        for idx, cell_val in enumerate(vals[:len(headers)]):
                            try:
                                row.cells[idx].text = cell_val
                            except Exception:
                                pass
                        # If fewer values than headers, remaining cells stay blank
            except Exception as ex:
                _log(f"[markdown_to_docx table parse] {repr(ex)}")
                for tl in table_lines:
                    add_paragraph_safe(doc, tl)

        # Bulleted list (-, *, +, •)
        elif re.match(r"^\s*[-*+] ", line) or re.match(r"^• ", line):
            items = []
            while i < len(lines) and (re.match(r"^\s*[-*+] ", lines[i]) or re.match(r"^• ", lines[i])):
                bullet = re.sub(r"^(\s*[-*+]|•)\s*", "", lines[i]).strip()
                items.append(bullet)
                i += 1
            i -= 1
            idx = 0
            while idx < len(items):
                bold_heading = re.match(r"^\*\*(.+?)\*\*:?$", items[idx])
                if bold_heading:
                    # Heading bullet
                    para = add_paragraph_safe(doc, "", style_name="List Bullet")
                    run = para.add_run(bold_heading.group(1) + (":" if items[idx].endswith(":") else ""))
                    run.bold = True
                    idx += 1
                    # Sub-items until next bold item
                    subidx = idx
                    while subidx < len(items) and not re.match(r"^\*\*.+?\*\*", items[subidx]):
                        # Try second-level bullet, fallback to first-level
                        try:
                            add_paragraph_safe(doc, items[subidx], style_name="List Bullet 2")
                        except Exception:
                            add_paragraph_safe(doc, items[subidx], style_name="List Bullet")
                        subidx += 1
                    idx = subidx
                else:
                    add_paragraph_safe(doc, items[idx], style_name="List Bullet")
                    idx += 1

        # Numbered list (1., 2., etc.)
        elif re.match(r"^\s*\d+\.", line):
            while i < len(lines) and re.match(r"^\s*\d+\.", lines[i]):
                txt = re.sub(r"^\s*\d+\.\s*", "", lines[i]).strip()
                try:
                    add_paragraph_safe(doc, txt, style_name="List Number")
                except Exception:
                    add_paragraph_safe(doc, txt)
                i += 1
            i -= 1

        # Horizontal rule --- ___ ***
        elif re.match(r"^(-{3,}|_{3,}|\*{3,})$", line.strip()):
            # python-docx has no HR; insert an empty paragraph for spacing
            add_paragraph_safe(doc, "")

        # Blockquote
        elif line.strip().startswith(">"):
            quote = line.lstrip("> ").strip()
            para = add_paragraph_safe(doc, "")
            run = para.add_run(quote)
            run.italic = True
            try:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            except Exception:
                pass

        # Normal paragraph
        elif line.strip() != "":
            add_paragraph_safe(doc, line)
        else:
            # Preserve blank lines with an empty paragraph
            add_paragraph_safe(doc, "")

        i += 1

    # Save document
    try:
        doc.save(docx_path)
    except Exception as ex:
        _log(f"[markdown_to_docx save] {repr(ex)}")
        raise


def robust_write_docx_file(md_text: str) -> str:
    import traceback

    os.makedirs(BASE_DIR, exist_ok=True)
    docx_name = f"chat_output_{uuid.uuid4().hex}.docx"
    docx_path = os.path.join(BASE_DIR, docx_name)
    try:
        if not isinstance(md_text, str):
            md_text = str(md_text)
        markdown_to_docx(md_text, docx_path)
        return f"/download_docx?file={docx_name}"
    except Exception:
        _log(f"[markdown_to_docx error]\n{traceback.format_exc()}")
        return ""