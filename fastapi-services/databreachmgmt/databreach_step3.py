# source venv/bin/activate
# /home/cybersecai/htdocs/www.cybersecai.io/webhook/venv/bin/gunicorn -w 1 -k uvicorn.workers.UvicornWorker -b 127.0.0.1:8302 docs_generator_agentic_ai:app --timeout 1200
#lsof -i :8302

# Agentic AI service for generating custom compliance documents as .docx files, using:

# LLM (via OpenAI API) for document text/content.
# Microsoft Word (.docx) output only (no PDF logic).
# Data per user, taken from user’s folder.
# Step-by-Step:
# Receives a POST request with:

# user_id (string, e.g. "1")
# json_data (list of sensitive data/classified info for that user)
# Looks in /home/cybersecai/htdocs/www.cybersecai.io/databreachmgmt/<user_id>/ for all .json files.

# Each JSON file should encode a policy/compliance document (or a set, depending on file structure).
# Loads all found JSONs, collects "documents" from each.

# Each document item is a template, which the LLM can rewrite/tailor.
# For each document:

# Asks the LLM (OpenAI API) whether this template can benefit from merging in classified data.
# If “yes”:
# Prompts the LLM to generate a real document in Markdown using the classified data.
# Converts this Markdown to .docx (MS Word file).
# Returns a JSON response listing the generated results and download links.

# Has endpoints for downloading generated .docx files.

import json
import logging
import os
import re
import shutil
import sys
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from openai import OpenAI
from pydantic import BaseModel

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
USER_DOC_DIR = "/home/cybersecai/htdocs/www.cybersecai.io/databreachmgmt"

LOG_PATH = "/tmp/agentic_service.log"
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout), logging.FileHandler(LOG_PATH, encoding="utf-8")]
)
logger = logging.getLogger("agentic_service")

app = FastAPI(title="Agentic AI Sensitive Document Generator")

# -------------------- Helper Models -------------------- #

class BusinessContext(BaseModel):
    industry: Optional[str] = ""
    country: Optional[str] = ""
    about_business: Optional[str] = ""

class DocumentRequest(BaseModel):
    user_id: str
    organisation_name: str
    json_data: List[Dict[str, Any]] = []
    business_context: Optional[BusinessContext] = None

# -------------------- Utility Functions -------------------- #

def detect_mandatory(doc: Dict[str, Any]) -> bool:
    text = str(doc.get("LegalOrBestPractice", "")).lower()
    return "mandatory" in text or "must" in text

def group_doc_type(doc_type: str) -> str:
    t = doc_type.lower()
    if "policy" in t:
        return "Policy"
    if "plan" in t:
        return "Plan"
    if "procedure" in t or "process" in t:
        return "Procedure"
    if any(word in t for word in ["register", "log", "record"]):
        return "Register/Log"
    return "Other"

def prettify_filename(s: str) -> str:
    base = os.path.basename(s)
    base = base.replace("_generated", "").replace(".docx", "").replace(".json", "")
    return base.replace("_", " ").title()

def safe_filename(text: str) -> str:
    return "".join(c if c.isalnum() or c in ("-", "_") else "_" for c in text)

def markdown_to_docx(markdown: str, output_path: str) -> None:
    doc = Document()
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if re.match(r"^#{1,6} ", line):
            level = len(line.split(" ")[0])
            doc.add_heading(line[level + 1:].strip(), min(level - 1, 4))
        elif re.match(r"^\s*[-*+]\s+", line) or line.strip().startswith("• "):
            bullet_items = []
            while i < len(lines) and (re.match(r"^\s*[-*+]\s+", lines[i]) or lines[i].strip().startswith("• ")):
                bullet = re.sub(r"^(\s*[-*+]|•)\s*", "", lines[i]).strip()
                bullet_items.append(bullet)
                i += 1
            i -= 1
            idx = 0
            while idx < len(bullet_items):
                bold_heading = re.match(r"^\*\*(.+?)\*\*:?$", bullet_items[idx])
                if bold_heading:
                    para = doc.add_paragraph(style="List Bullet")
                    run = para.add_run(bold_heading.group(1) + (":" if bullet_items[idx].endswith(":") else ""))
                    run.bold = True
                    idx += 1
                    while idx < len(bullet_items) and not re.match(r"^\*\*.+?\*\*", bullet_items[idx]):
                        doc.add_paragraph(bullet_items[idx], style="List Bullet 2")
                        idx += 1
                else:
                    doc.add_paragraph(bullet_items[idx], style="List Bullet")
                    idx += 1
        elif re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                txt = re.sub(r"^\s*\d+\.\s*", "", lines[i]).strip()
                doc.add_paragraph(txt, style="List Number")
                i += 1
            i -= 1
        elif re.match(r"^(-{3,}|_{3,}|\\*{3,})$", line.strip()):
            doc.add_paragraph().add_run().add_break()
        elif line.strip().startswith(">"):
            quote = line.lstrip("> ").strip()
            para = doc.add_paragraph()
            run = para.add_run(quote)
            run.italic = True
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.strip():
            doc.add_paragraph(line)
        i += 1
    doc.save(output_path)

def ask_benefit(doc_type: str, template: Dict[str, Any], sensitive_data: List[Dict[str, Any]], context: Dict[str, Any]) -> Dict[str, Any]:
    summary = template.get("BriefDescription", "")
    ctx_text = f"Industry: {context.get('industry','')}, Country: {context.get('country','')}"
    prompt = f"""
Document Type: {doc_type}
Template summary: {summary}
Business context: {ctx_text}
High-risk snippets: {json.dumps(sensitive_data)[:2000] or 'None'}

Should we enrich this document with the snippets above? Respond ONLY with JSON:
{{
  "can_benefit": true/false,
  "reason": "short explanation"
}}
"""
    try:
        resp = client.chat.completions.create(
            model="gpt-4.1",
            temperature=0.1,
            messages=[
                {"role": "system", "content": "You are a pragmatic CISO/legal expert."},
                {"role": "user", "content": prompt}
            ]
        )
        return json.loads(resp.choices[0].message.content)
    except Exception as exc:
        logger.error(f"Benefit check failed: {exc}")
        return {"can_benefit": bool(sensitive_data), "reason": "Fallback decision"}

def generate_markdown(doc_type: str, template: Dict[str, Any], sensitive_data: List[Dict[str, Any]], context: Dict[str, Any]) -> str:
    prompt = f"""
Organisation context:
- Industry: {context.get('industry','')}
- Country: {context.get('country','')}
- About: {context.get('about_business','') or 'N/A'}

Document type: {doc_type}
Base template:
{template.get('DocumentTemplateContent','')[:10000]}

Sensitive data highlights (from cybersecai platform):
{json.dumps(sensitive_data, ensure_ascii=False)[:10000]}

Produce a polished Markdown document referencing the sensitive insights only where relevant. Include branded headers, RACI/accountability, triggers, communications guidance and audit considerations.
"""
    try:
        completion = client.chat.completions.create(
            model="gpt-4.1",
            temperature=0.35,
            messages=[
                {"role": "system", "content": "You are a senior incident-response lawyer."},
                {"role": "user", "content": prompt}
            ]
        )
        output = completion.choices[0].message.content.strip()
        if output.startswith("```"):
            output = output.strip("` \n")
            output = re.sub(r"^markdown\n", "", output, flags=re.I)
        return output
    except Exception as exc:
        logger.error(f"Markdown generation failed: {exc}")
        return "# ERROR\nUnable to generate document."

def load_templates_for_user(user_id: str) -> Dict[str, Any]:
    folder = f"{USER_DOC_DIR}/{user_id}"
    processed_folder = f"{folder}/processed"
    os.makedirs(processed_folder, exist_ok=True)

    docs = []
    doc_file_map = {}
    if not os.path.isdir(folder):
        raise HTTPException(404, "Template directory not found.")

    for file in Path(folder).glob("*.json"):
        path = str(file)
        if "generated" in path:
            continue

        with open(path, "r") as f:
            payload = json.load(f)

        doc_list = []
        if isinstance(payload, dict) and "documents" in payload:
            doc_list = payload["documents"]
        elif isinstance(payload, list):
            doc_list = payload
        elif isinstance(payload, dict):
            doc_list = [payload]

        for doc in doc_list:
            docs.append(doc)
            doc_file_map[doc.get("DocumentType")] = path

        try:
            shutil.move(path, f"{processed_folder}/{file.name}")
        except Exception as exc:
            logger.warning(f"Unable to move processed file {path}: {exc}")

    if not docs:
        raise HTTPException(status_code=404, detail="No templates available.")
    return {"docs": docs, "map": doc_file_map}

# -------------------- API Endpoints -------------------- #

@app.post("/agentic/generate_sensitive_docs")
def generate_sensitive_docs(request: DocumentRequest):
    logger.info(f"Starting sensitive doc generation for user {request.user_id}")
    template_bundle = load_templates_for_user(request.user_id)
    docs = template_bundle["docs"]
    doc_file_map = template_bundle["map"]
    context = (request.business_context or BusinessContext()).dict()

    user_folder = os.path.join(USER_DOC_DIR, request.user_id)
    os.makedirs(user_folder, exist_ok=True)

    results = []
    for index, doc in enumerate(docs, start=1):
        doc_type = doc.get("DocumentType", f"Doc_{uuid.uuid4().hex[:6]}")
        logger.info(f"[{index}/{len(docs)}] Processing {doc_type}")

        benefit = ask_benefit(doc_type, doc, request.json_data, context)
        if not benefit.get("can_benefit"):
            logger.info(f"Skipping {doc_type} ({benefit.get('reason')})")
            continue

        markdown = generate_markdown(doc_type, doc, request.json_data, context)

        source_file = doc_file_map.get(doc_type)
        base_dir = os.path.dirname(source_file) if source_file else user_folder
        generated_dir = os.path.join(base_dir, "generated")
        os.makedirs(generated_dir, exist_ok=True)

        safe_name = safe_filename(doc_type)
        json_path = os.path.join(generated_dir, f"{safe_name}_generated.json")
        docx_path = os.path.join(generated_dir, f"{safe_name}_generated.docx")

        enriched = dict(doc)
        enriched.update({
            "DocumentType": doc_type,
            "markdown": markdown,
            "benefit_reason": benefit.get("reason"),
            "organisation_name": request.organisation_name,
            "business_context": context,
        })

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, indent=2)

        markdown_to_docx(markdown, docx_path)

        results.append({
            **enriched,
            "doc_group": group_doc_type(doc_type),
            "is_mandatory": detect_mandatory(doc),
            "output_json": json_path,
            "output_docx": docx_path,
            "file_display_name": prettify_filename(docx_path),
        })

    if not results:
        raise HTTPException(status_code=200, detail="No relevant documents created.")

    return {"status": "ok", "documents": results, "user_folder": user_folder}

@app.get("/agentic/download_json/{user_id}/{filename:path}")
def download_json(user_id: str, filename: str):
    base = os.path.join(USER_DOC_DIR, user_id)
    target = os.path.join(base, filename)
    if os.path.isfile(target):
        return FileResponse(target, media_type="application/json", filename=os.path.basename(filename))
    for path in Path(base).rglob("generated"):
        candidate = os.path.join(path, os.path.basename(filename))
        if os.path.isfile(candidate):
            return FileResponse(candidate, media_type="application/json", filename=os.path.basename(filename))
    raise HTTPException(404, "File not found.")

@app.get("/agentic/download_docx/{user_id}/{filename:path}")
def download_docx(user_id: str, filename: str):
    base = os.path.join(USER_DOC_DIR, user_id)
    target = os.path.join(base, filename)
    if os.path.isfile(target):
        return FileResponse(target, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=os.path.basename(filename))
    for path in Path(base).rglob("generated"):
        candidate = os.path.join(path, os.path.basename(filename))
        if os.path.isfile(candidate):
            return FileResponse(candidate, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=os.path.basename(filename))
    raise HTTPException(404, "File not found.")

@app.get("/agentic/healthz")
def healthcheck():
    return {"status": "ok"}