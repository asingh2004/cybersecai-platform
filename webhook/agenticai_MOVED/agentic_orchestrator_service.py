#sudo lsof -i :8223
#sudo systemctl restart agentic_orchestrator_service
#sudo systemctl status agentic_orchestrator_service


import os, json, re, glob
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from datetime import datetime, timedelta
import openai
import uuid
import dateutil.parser
from uuid import uuid4
from collections import Counter, defaultdict
from typing import List, Dict, Any, Optional, Tuple, Union
from fastapi.responses import FileResponse
import csv
import markdown2
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



# CONFIG
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY environment variable not set.")
client = openai.OpenAI(api_key=OPENAI_API_KEY)
LARAVEL_LOG_PATH = '/home/cybersecai/htdocs/www.cybersecai.io/storage/logs/laravel.log'

app = FastAPI()
session_mind = {}

# LARAVEL LOGGING
def log_to_laravel(message: str):
    timestamp = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
    with open(LARAVEL_LOG_PATH, 'a', encoding='utf-8') as log_file:
        log_file.write(f"[{timestamp}] python.INFO: {message}\n")

# ===== UTILITIES (unaltered from your original code) =====
def output_guardrails(text: str) -> None:
    bad_words = ['racist', 'intercourse']
    if any(re.search(rf"\b{w}\b", text, re.I) for w in bad_words):
        raise HTTPException(400, "Inappropriate content detected in LLM Output")
    if len(text) > 30000:
        raise HTTPException(400, "Output exceeds max length")


def format_high_risk_files_markdown(findings):
    out = []
    out.append("### All High Risk Files: Detailed Inventory\n")
    if not findings:
        out.append("_No high risk files found._")
        return "\n".join(out)
    for f in findings:
        out.append("---")
        out.append(f"#### **{f.get('file_name', '[No Name]')}**\n")
        out.append("**File Details:**")
        out.append(f"- **File Name:** {f.get('file_name','')}")
        out.append(f"- **Full Path:** {f.get('full_path','')}")
        out.append(f"- **Last Modified:** {f.get('last_modified','')}")
        out.append(f"- **Created:** {f.get('created','') or '_None_'}")
        out.append(f"- **Data Source:** {f.get('data_source','')}")
        out.append(f"- **Classification:** {f.get('data_classification','')}")
        out.append(f"- **Data Subject Area:** {f.get('likely_data_subject_area','')}")
        out.append(f"- **Overall Risk Rating:** {f.get('overall_risk_rating','')}")

        auditor_view = f.get("auditor_agent_view")
        if auditor_view:
            out.append(f"**Auditor View:**  \n> {auditor_view}")

        auditor_action = f.get("auditor_proposed_action")
        if auditor_action:
            out.append(f"**Auditor Proposed Action:**  \n{auditor_action}")

        controls = f.get("cyber_proposed_controls")
        if controls:
            out.append("**Cyber Proposed Controls:**")
            if isinstance(controls, list):
                for c in controls:
                    out.append(f"- {c}")
            else:
                out.append(f"{controls}")

        # Compliance Findings Table
        cf = f.get('compliance_findings', [])
        if cf:
            out.append("\n**Compliance Findings:**")
            out.append("| Standard | Jurisdiction | Detected Fields | Risk |")
            out.append("|----------|--------------|-----------------|------|")
            for c in cf:
                out.append(
                    f"| {c.get('standard','')} | {c.get('jurisdiction','')} | {', '.join(c.get('detected_fields', []) or [])} | {c.get('risk','')} |"
                )
        out.append("")  # Space after each file
    return "\n".join(out)


def format_all_risks_files_markdown(findings):
    out = []
    out.append("### All Risk Files: Detailed Inventory\n")
    if not findings:
        out.append("_No risk files found._")
        return "\n".join(out)
    for f in findings:
        # If you only want files with ANY risk, not "None" or blank:
        risk = (f.get('overall_risk_rating','') or '').lower()
        if risk in ['none', '', None]:
            continue
        out.append("---")
        out.append(f"#### **{f.get('file_name', '[No Name]')}**\n")
        out.append("**File Details:**")
        out.append(f"- **File Name:** {f.get('file_name','')}")
        out.append(f"- **Full Path:** {f.get('full_path','')}")
        out.append(f"- **Last Modified:** {f.get('last_modified','')}")
        out.append(f"- **Created:** {f.get('created','') or '_None_'}")
        out.append(f"- **Data Source:** {f.get('data_source','')}")
        out.append(f"- **Classification:** {f.get('data_classification','')}")
        out.append(f"- **Data Subject Area:** {f.get('likely_data_subject_area','')}")
        out.append(f"- **Overall Risk Rating:** {f.get('overall_risk_rating','')}")
        auditor_view = f.get("auditor_agent_view")
        if auditor_view:
            out.append(f"**Auditor View:**  \n> {auditor_view}")
        auditor_action = f.get("auditor_proposed_action")
        if auditor_action:
            out.append(f"**Auditor Proposed Action:**  \n{auditor_action}")
        controls = f.get("cyber_proposed_controls")
        if controls:
            out.append("**Cyber Proposed Controls:**")
            if isinstance(controls, list):
                for c in controls:
                    out.append(f"- {c}")
            else:
                out.append(f"{controls}")
        # Compliance Findings Table
        cf = f.get('compliance_findings', [])
        if cf:
            out.append("\n**Compliance Findings:**")
            out.append("| Standard | Jurisdiction | Detected Fields | Risk |")
            out.append("|----------|--------------|-----------------|------|")
            for c in cf:
                out.append(
                    f"| {c.get('standard','')} | {c.get('jurisdiction','')} | {', '.join(c.get('detected_fields', []) or [])} | {c.get('risk','')} |"
                )
        out.append("")  # Space after each file
    if len(out) == 1:
        out.append("_No risk files found._") # Only header; i.e. all were none risk.
    return "\n".join(out)


def format_medium_risk_files_markdown(findings):
    out = []
    out.append("### All Medium Risk Files: Detailed Inventory\n")
    found_any = False
    for f in findings:
        risk = (f.get('overall_risk_rating','') or '').lower()
        if risk != 'medium':
            continue
        found_any = True
        out.append("---")
        out.append(f"#### **{f.get('file_name', '[No Name]')}**\n")
        out.append("**File Details:**")
        out.append(f"- **File Name:** {f.get('file_name','')}")
        out.append(f"- **Full Path:** {f.get('full_path','')}")
        out.append(f"- **Last Modified:** {f.get('last_modified','')}")
        out.append(f"- **Created:** {f.get('created','') or '_None_'}")
        out.append(f"- **Data Source:** {f.get('data_source','')}")
        out.append(f"- **Classification:** {f.get('data_classification','')}")
        out.append(f"- **Data Subject Area:** {f.get('likely_data_subject_area','')}")
        out.append(f"- **Overall Risk Rating:** {f.get('overall_risk_rating','')}")
        auditor_view = f.get("auditor_agent_view")
        if auditor_view:
            out.append(f"**Auditor View:**  \n> {auditor_view}")
        auditor_action = f.get("auditor_proposed_action")
        if auditor_action:
            out.append(f"**Auditor Proposed Action:**  \n{auditor_action}")
        controls = f.get("cyber_proposed_controls")
        if controls:
            out.append("**Cyber Proposed Controls:**")
            if isinstance(controls, list):
                for c in controls:
                    out.append(f"- {c}")
            else:
                out.append(f"{controls}")
        # Compliance Findings Table
        cf = f.get('compliance_findings', [])
        if cf:
            out.append("\n**Compliance Findings:**")
            out.append("| Standard | Jurisdiction | Detected Fields | Risk |")
            out.append("|----------|--------------|-----------------|------|")
            for c in cf:
                out.append(
                    f"| {c.get('standard','')} | {c.get('jurisdiction','')} | {', '.join(c.get('detected_fields', []) or [])} | {c.get('risk','')} |"
                )
        out.append("")  # Space after each file
    if not found_any:
        out.append("_No medium risk files found._")
    return "\n".join(out)


def agent_findings_facts(data: Dict[str, Any]):
    import json
    op = data.get("operation")
    args = data.get("args", {}) if data.get("args") is not None else {}
    config_ids = data.get("config_ids")
    user_query = (data.get("query") or op or "").lower()
    
    # Use gather_all_findings for any non-high risk or all risk queries
    if any(kw in user_query for kw in ["medium risk", "low risk", "all risk", "any risk"]):
        findings = gather_all_findings(config_ids)
    else:
        findings = gather_high_risk_findings(config_ids)

    CSV_FIELD_MAP = {
        "data subject area": "likely_data_subject_area",
        "subject area": "likely_data_subject_area",
        "data classification": "data_classification",
        "classification": "data_classification",
        "risk rating": "overall_risk_rating",
        "file name": "file_name",
        "file path": "file_path",
        "full path": "full_path",
        "last modified": "last_modified",
        "created": "created",
        "permissions": "permissions",
        "data source": "data_source"
    }
    COMPLIANCE_FIELD_MAP = {
        "standard": "standard",
        "jurisdiction": "jurisdiction",
        "detected field": "detected_fields",
        "detected fields": "detected_fields",
        "finding risk": "risk"
    }

    def md_table(d, title=""):
        if not d:
            return f"_No results found._"
        headers = ["Item", "Count"]
        lines = []
        if title:
            lines.append(f"### {title}")
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|------|-------|")
        for k, v in d.items():
            lines.append(f"| {k} | {v} |")
        return "\n".join(lines)

    def md_key_list(values, key):
        if not values:
            return f"_No {key} values found in current evidence._"
        msg = f"# Unique {key.title()}\n"
        msg += "\n".join(f"- {v}" for v in sorted(values))
        return msg

    # --- NICE MARKDOWN OUTPUT FOR RISK QUERIES ---
    if re.search(r"(how many|count|number of).*(file|files).*(medium risk)", user_query):
        n = sum(1 for f in findings if (f.get('overall_risk_rating') or '').strip().lower() == 'medium')
        reply = f"### Medium Risk File Summary\n\n| Overall Risk Rating | Number of Files |\n|--------------------|----------------|\n| Medium | {n} |\n"
        #return {"reply": reply, "raw": n, "operation": op, "args": args}
        return reply

    if re.search(r"(how many|count|number of).*(file|files).*(low risk)", user_query):
        n = sum(1 for f in findings if (f.get('overall_risk_rating') or '').strip().lower() == 'low')
        reply = f"### Low Risk File Summary\n\n| Overall Risk Rating | Number of Files |\n|--------------------|----------------|\n| Low | {n} |\n"
        #return {"reply": reply, "raw": n, "operation": op, "args": args}
        return reply

    if re.search(r"(how many|count|number of).*(file|files).*(risk)", user_query):
        risk_counts = {}
        for f in findings:
            rating = (f.get('overall_risk_rating') or '').strip().capitalize()
            if rating and rating.lower() not in ("", "none"):
                risk_counts[rating] = risk_counts.get(rating, 0) + 1
        if risk_counts:
            lines = ["### File Count by Risk Rating",
                     "| Overall Risk Rating | Number of Files |",
                     "|--------------------|----------------|"]
            for k, v in sorted(risk_counts.items()):
                lines.append(f"| {k} | {v} |")
            reply = "\n".join(lines)
        else:
            reply = "_No files with a risk rating detected._"
        return {"reply": reply, "raw": risk_counts, "operation": op, "args": args}
    # ------------------------------------------------

    # 1. "group_by_field" operation, robust for both csv and compliance_findings
    if op == "group_by_field":
        field = args.get("field")
        if not field:
            return {"reply": "_No field specified for group_by_field operation._", "raw": {}, "operation": op, "args": args}
        # Compliance fields
        if field in COMPLIANCE_FIELD_MAP.values():
            from collections import Counter
            c = Counter()
            for f in findings:
                compliance = f.get("compliance_findings", [])
                if compliance and isinstance(compliance, list):
                    for r in compliance:
                        val = r.get(field)
                        # Detected fields (list-of)
                        if isinstance(val, list):
                            for v in val:
                                if v: c[v] += 1
                        elif val:
                            c[str(val)] += 1
            result = {k: v for k, v in c.items() if k and k.lower() not in ('[missing]', '[none]', 'none')}
            if not result:
                result = {"[None found]": 0}
            reply = md_table(result, title=f"Files by {field.title()}")
            #return {"reply": reply, "raw": result, "operation": op, "args": args}
            return reply
        # Top-level fields
        else:
            from collections import Counter
            c = Counter(str(f.get(field, '[Missing]')) for f in findings)
            result = {k: v for k, v in c.items() if k and k.lower() not in ('[missing]', '[none]', 'none')}
            if not result:
                result = {"[None found]": 0}
            reply = md_table(result, title=f"Files by {field.title()}")
            #return {"reply": reply, "raw": result, "operation": op, "args": args}
            return reply

    # 2. "list unique X" for all fields (top-level and compliance)
    for key, field in {**CSV_FIELD_MAP, **COMPLIANCE_FIELD_MAP}.items():
        patterns = [
            fr"(list|show|unique|all).*{re.escape(key)}",
            fr"{re.escape(key)}.*(list|show|unique|all)"
        ]
        if op and re.search(fr"{re.escape(field)}", op, re.I):
            patterns = ["^.*$"]  # force match if op matches field
        if any(re.search(p, user_query) for p in patterns) or (op and op.startswith("unique_") and field in op):
            # Top-level field
            if field in CSV_FIELD_MAP.values():
                unique = sorted(set(str(f.get(field, '')).strip() for f in findings if f.get(field)))
            else:
                unique = set()
                for f in findings:
                    compliance = f.get("compliance_findings", [])
                    if compliance and isinstance(compliance, list):
                        for r in compliance:
                            val = r.get(field)
                            if isinstance(val, list):
                                for v in val:
                                    if v: unique.add(v)
                            elif val:
                                unique.add(str(val))
                unique = sorted(unique)
            reply = md_key_list(unique, key)
            return {"reply": reply, "raw": unique, "operation": op, "args": args}

    # 3. Pretty Table for files_by_source
    if op == "files_by_source":
        ds = args.get("data_source", "")
        files = [f for f in findings if f.get("data_source", "").lower() == ds.lower()]
        if files:
            def flatten(row):
                row_flat = dict(row)
                if "compliance_findings" in row_flat:
                    row_flat["found_standards"] = ", ".join(
                        str(r.get("standard", "")) for r in row_flat["compliance_findings"] if r.get("standard"))
                return row_flat
            headers = list(flatten(files[0]).keys())
            lines = [
                "File Listing:",
                "| " + " | ".join(headers) + " |",
                "| " + " | ".join(["---"] * len(headers)) + " |"
            ]
            for f in files:
                flat = flatten(f)
                lines.append("| " + " | ".join(str(flat.get(h, "")) for h in headers) + " |")
            reply = "\n".join(lines)
        else:
            reply = "_No files from this data source found._"
        return {"reply": reply, "raw": files, "operation": op, "args": args}

    # 4. Fallback: pretty JSON dump
    #pretty = json.dumps(findings, indent=2, ensure_ascii=False)
    #reply = f"```json\n{pretty}\n```"
    #return {"reply": reply, "raw": findings, "operation": op, "args": args}
    pretty = json.dumps(findings, indent=2, ensure_ascii=False)
    reply = f"```json\n{pretty}\n```"
    return reply

def gather_all_findings(config_ids, base_dir="/home/cybersecai/htdocs/www.cybersecai.io/webhook", base_paths=None):
    if base_paths is None:
        base_paths = ['M365', 'SMB', 'S3', 'NFS']
    findings = []
    for cid in config_ids:
        for backend in base_paths:
            graph_dir = os.path.join(base_dir, backend, str(cid), "graph")
            if not os.path.isdir(graph_dir):
                continue
            for filepath in glob.glob(os.path.join(graph_dir, "*.json")):
                try:
                    with open(filepath, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        items = data if isinstance(data, list) else [data]
                        for item in items:
                            if not isinstance(item, dict):
                                continue
                            llm = item.get('llm_response')
                            if not llm:
                                continue
                            try:
                                if isinstance(llm, str):
                                    if llm.startswith('"') and llm.endswith('"'):
                                        llm = llm[1:-1]
                                    llm = json.loads(llm)
                            except Exception:
                                continue
                            compliance_findings = []
                            results = llm.get('results', [])
                            if isinstance(results, list):
                                for r in results:
                                    if not isinstance(r, dict):
                                        continue
                                    compliance_findings.append({
                                        'standard': r.get('standard'),
                                        'jurisdiction': r.get('jurisdiction'),
                                        'detected_fields': r.get('detected_fields', []),
                                        'risk': r.get('risk')
                                    })
                            permissions = item.get('permissions', None)
                            finding = {
                                'file_name': item.get('file_name'),
                                'last_modified': item.get('last_modified'),
                                'created': item.get('created'),
                                'file_path': item.get('file_path'),
                                'full_path': item.get('full_path'),
                                'compliance_findings': compliance_findings,
                                'auditor_agent_view': llm.get('auditor_agent_view'),
                                'data_classification': llm.get('data_classification'),
                                'likely_data_subject_area': llm.get('likely_data_subject_area'),
                                'overall_risk_rating': llm.get('overall_risk_rating'),
                                'cyber_proposed_controls': llm.get('cyber_proposed_controls'),
                                'auditor_proposed_action': llm.get('auditor_proposed_action'),
                                'permissions': permissions,
                                'data_source': backend
                            }
                            findings.append(finding)
                except Exception as ex:
                    log_to_laravel(f"Error reading {filepath}: {repr(ex)}")
                    continue
    return findings


## The function below only Yields or collect only files with overall_risk_rating == 'High'
def gather_high_risk_findings(config_ids, base_dir="/home/cybersecai/htdocs/www.cybersecai.io/webhook", base_paths=None):
    """
    Scans each config/connection/source for risk analysis result .json files.
    Extracts exactly ONE normalized finding per file, but ONLY includes files with overall_risk_rating == 'High'.
    Returns a list of dicts, one per file.
    """
    if base_paths is None:
        base_paths = ['M365', 'SMB', 'S3', 'NFS']
    findings = []

    for cid in config_ids:
        for backend in base_paths:
            graph_dir = os.path.join(base_dir, backend, str(cid), "graph")
            if not os.path.isdir(graph_dir):
                continue
            for filepath in glob.glob(os.path.join(graph_dir, "*.json")):
                try:
                    with open(filepath, "r", encoding="utf-8") as f:
                        content = f.read()
                        try:
                            data = json.loads(content)
                        except Exception:
                            continue
                        items = data if isinstance(data, list) else [data]
                        for item in items:
                            if not isinstance(item, dict):
                                continue
                            llm = item.get('llm_response')
                            if not llm:
                                continue
                            try:
                                if isinstance(llm, str):
                                    if llm.startswith('"') and llm.endswith('"'):
                                        llm = llm[1:-1]
                                    llm = json.loads(llm)
                            except Exception:
                                continue

                            if not isinstance(llm, dict):
                                continue

                            # ---- Filter for overall_risk_rating == 'High' only ----
                            risk_rating = llm.get('overall_risk_rating')
                            # Case-insensitive, tolerate whitespace, support integer-like as well
                            if not risk_rating or str(risk_rating).strip().lower() != "high":
                                continue

                            # Extract compliance findings
                            compliance_findings = []
                            results = llm.get('results', [])
                            if isinstance(results, list):
                                for r in results:
                                    if not isinstance(r, dict):
                                        continue
                                    compliance_findings.append({
                                        'standard': r.get('standard'),
                                        'jurisdiction': r.get('jurisdiction'),
                                        'detected_fields': r.get('detected_fields', []),
                                        'risk': r.get('risk')
                                    })

                            permissions = item.get('permissions', None)

                            finding = {
                                'file_name': item.get('file_name'),
                                'last_modified': item.get('last_modified'),
                                'created': item.get('created'),
                                'file_path': item.get('file_path'),
                                'full_path': item.get('full_path'),
                                'compliance_findings': compliance_findings,
                                'auditor_agent_view': llm.get('auditor_agent_view'),
                                'data_classification': llm.get('data_classification'),
                                'likely_data_subject_area': llm.get('likely_data_subject_area'),
                                'overall_risk_rating': risk_rating,
                                'cyber_proposed_controls': llm.get('cyber_proposed_controls'),
                                'auditor_proposed_action': llm.get('auditor_proposed_action'),
                                'permissions': permissions,
                                'data_source': backend
                            }
                            findings.append(finding)
                except Exception:
                    # For production, consider logging errors (omitted for clarity)
                    log_to_laravel(f"Error reading {filepath}: {repr(ex)}")
                    continue
    return findings


def parse_date_from_query(query):
    m = re.search(r'last\s+(\d+)\s+day', query, re.I)
    if m:
        n = int(m.group(1))
        end = datetime.utcnow()
        start = end - timedelta(days=n)
        return (start, end)
    if re.search(r'(today|current day)', query, re.I):
        now = datetime.utcnow()
        return (now.replace(hour=0, minute=0, second=0, microsecond=0),
                now.replace(hour=23, minute=59, second=59, microsecond=999999))
    m = re.search(r'(\d{1,2}\s+\w+\s+\d{4})', query)
    if m:
        try:
            d = dateutil.parser.parse(m.group(1), fuzzy=True).date()
            start = datetime.combine(d, datetime.min.time())
            end = datetime.combine(d, datetime.max.time())
            return (start, end)
        except Exception:
            pass
    return None




# ===== AGENT LOGIC: ALL PROMPTS COPIED FROM YOUR CODE, as-is =====

class ChatBotReq(BaseModel):
    persona: str
    query: str
    messages: list
    config_ids: list

def agent_chatbot(data: Dict[str, Any]):
    class DummyChatReq:
        persona: str
        query: str
        messages: list
        config_ids: list
        def __init__(self, persona, query, messages, config_ids):
            self.persona = persona
            self.query = query
            self.messages = messages
            self.config_ids = config_ids
    # retain identical prompt logic
    cbreq = DummyChatReq(
        data['persona'],
        data['query'],
        data['messages'],
        data['config_ids']
    )
    # direct call!
    return agent_chatbot_actual(cbreq)


def agent_chatbot_actual(req: ChatBotReq):

    # Map user language to top-level CSV fields (update as new fields added)
    CSV_FIELD_MAP = {
        "data subject area": "likely_data_subject_area",
        "subject area": "likely_data_subject_area",
        "data classification": "data_classification",
        "classification": "data_classification",
        "risk rating": "overall_risk_rating",
        "file name": "file_name",
        "file path": "file_path",
        "full path": "full_path",
        "last modified": "last_modified",
        "created": "created",
        "permissions": "permissions",
        "auditor agent view": "auditor_agent_view",
        "auditor proposed action": "auditor_proposed_action",
        "Immediate Containment": "auditor_proposed_action",
        "Containment": "auditor_proposed_action",
        "auditor agent insight": "auditor_agent_view",
        "recommended action": "auditor_proposed_action",
        "cyber proposed controls": "cyber_proposed_controls",
        "Compliance and Notification": "Cybersecai Data Breach AI AGent can expertly guide you through the process and deetrmination on notification. including generation of any communication letters etc.",
        "Compliance": "Cybersecai Data Breach AI AGent can expertly guide you through the process and deetrmination on notification. including generation of any communication letters etc.",
        "Notification": "Cybersecai Data Breach AI AGent can expertly guide you through the process and deetrmination on notification. including generation of any communication letters etc.",
        "Data Handling and Remediation": "cyber_proposed_controls",
        "Ongoing Monitoring and Prevention": "You are in safe hands, as vybersecai.io platform coducts ongoing minority of all files",

        # "compliance_findings" is handled separately!
    }

    # Map queries to compliance_findings subfields
    COMPLIANCE_FIELD_MAP = {
        "standard": "standard",
        "jurisdiction": "jurisdiction",
        "detected field": "detected_fields",  # This is a list
        "detected fields": "detected_fields",
        "finding risk": "risk"
    }

    # --- Begin logic ---
    user_query = (req.query or "").lower().strip()
    #findings = gather_high_risk_findings(req.config_ids)

    # Smart loader: use all for medium, low, or "any"; filter in code!
    user_query = (req.query if hasattr(req, 'query') else data.get("query", "")).lower()

    # Use all findings when not exclusive to high risk
    if any(kw in user_query for kw in ["medium risk", "low risk", "all risk", "any risk"]):
        findings = gather_all_findings(req.config_ids)
    else:
        findings = gather_high_risk_findings(req.config_ids)



    # List all details for every file currently rated high risk
    # if re.search(r'(list|show|all).*(detail|record|full|everything)(.*high risk)?', user_query):
    #     markdown = format_high_risk_files_markdown(findings)
    #     return {"reply": markdown, "raw": findings}


    # List all details for every file currently rated high risk
    if re.search(r'(list|show|all).*(detail|record|full|everything).*(high risk)', user_query, re.I):
        markdown = format_high_risk_files_markdown(findings)
        return {"reply": markdown, "raw": findings}

    # List all details for every file currently rated medium risk
    elif re.search(r'(list|show|all).*(detail|record|full|everything).*(medium risk)', user_query, re.I):
        markdown = format_medium_risk_files_markdown(findings)
        return {"reply": markdown, "raw": findings}

    # List all details for every file currently rated any risk (excluding 'none')
    elif re.search(r'(list|show|all).*(detail|record|full|everything).*(risk)', user_query, re.I):
        markdown = format_all_risks_files_markdown(findings)
        return {"reply": markdown, "raw": findings}

    # --- Data classification count summary ---
    if re.search(r"(how many|count|number of).*(file|files).*(data classification|classification)", user_query):
        class_count = {}
        for f in findings:
            cls = (f.get("data_classification") or "").strip()
            if cls:
                class_count[cls] = class_count.get(cls, 0) + 1
            else:
                class_count["[Unclassified]"] = class_count.get("[Unclassified]", 0) + 1
        if not class_count:
            reply = "_No data classifications found in current high risk files._"
        else:
            reply = "### High Risk File Count by Data Classification\n\n"
            reply += "| Data Classification | File Count |\n"
            reply += "|---------------------|------------|\n"
            for k, v in class_count.items():
                reply += f"| {k} | {v} |\n"
        return {"reply": reply, "raw": class_count}



        # --- Data classification count summary ---
    if re.search(r"(how many|count|number of).*(file|files).*(data classification|classification)", user_query):
        class_count = {}
        for f in findings:
            cls = (f.get("data_classification") or "").strip()
            if cls:
                class_count[cls] = class_count.get(cls, 0) + 1
            else:
                class_count["[Unclassified]"] = class_count.get("[Unclassified]", 0) + 1
        if not class_count:
            reply = "_No data classifications found in current high risk files._"
        else:
            reply = "### High Risk File Count by Data Classification\n\n"
            reply += "| Data Classification | File Count |\n"
            reply += "|---------------------|------------|\n"
            for k, v in class_count.items():
                reply += f"| {k} | {v} |\n"
        return {"reply": reply, "raw": class_count}

    # 1. Direct display from CSV top-level fields
    for key, field in CSV_FIELD_MAP.items():
        patterns = [
            fr"(list|show|unique|all).*{re.escape(key)}",
            fr"{re.escape(key)}.*(list|show|unique|all)"
        ]
        if any(re.search(p, user_query) for p in patterns):
            unique_values = sorted(set(
                str(f.get(field, '')).strip()
                for f in findings if f.get(field)
            ))
            # For permissions (dict/list), flatten as needed:
            if field == "permissions":
                perms = set()
                for f in findings:
                    pval = f.get("permissions")
                    if not pval:
                        continue
                    if isinstance(pval, str):
                        perms.add(pval)
                    elif isinstance(pval, list):
                        for p in pval:
                            perms.add(str(p))
                    elif isinstance(pval, dict):
                        perms.add(str(pval))
                unique_values = sorted(perms)
            if unique_values:
                reply = f"# Unique {key.title()}\n" + "\n".join(f"- {x}" for x in unique_values)
            else:
                reply = f"_No {key} values found in the current evidence._"
            return {"reply": reply}

    # 2. Direct display of compliance_findings subfields (e.g. list all detected standards/jurisdictions)
    for key, field in COMPLIANCE_FIELD_MAP.items():
        patterns = [
            fr"(list|show|unique|all).*{re.escape(key)}",
            fr"{re.escape(key)}.*(list|show|unique|all)"
        ]
        if any(re.search(p, user_query) for p in patterns):
            results_set = set()
            for f in findings:
                compliance = f.get("compliance_findings", [])
                if not compliance or not isinstance(compliance, list):
                    continue
                for r in compliance:
                    if field == "detected_fields" and isinstance(r.get(field), list):
                        for val in r[field]:
                            if val: results_set.add(val)
                    else:
                        value = r.get(field)
                        if value: results_set.add(str(value))
            unique_values = sorted(results_set)
            if unique_values:
                reply = f"# Unique Compliance Finding {key.title()}\n" + "\n".join(f"- {v}" for v in unique_values)
            else:
                reply = f"_No {key} values found in compliance findings._"
            return {"reply": reply}

    # --- Summarize available files by storage backend (data_source) ---
    storage_count = {}
    backend_files = {}
    for entry in findings:
        ds = (entry.get('data_source') or "").upper()
        if ds:
            storage_count[ds] = storage_count.get(ds, 0) + 1
            backend_files.setdefault(ds, []).append(entry.get("file_name", ""))
    # Build a Markdown table and bullet list for the LLM prompt
    storage_md = ""
    if storage_count:
        storage_md = "### File Count by Storage Backend\n\n"
        storage_md += "| Storage Backend | Files Found |\n"
        storage_md += "|----------------|-------------|\n"
        for ds, cnt in storage_count.items():
            storage_md += f"| {ds} | {cnt} |\n"
        storage_md += "\n"
        for ds, files in backend_files.items():
            storage_md += f"**{ds} Files:**\n"
            for fname in files:
                storage_md += f"- {fname}\n"
            storage_md += "\n"
    else:
        storage_md = "_No high risk files found in any backend._\n"


    # 3. All files listing (optional)
    if re.search(r'list all files|show all files|all file names', user_query):
        all_files = sorted(set(str(f.get("file_name", "")).strip() for f in findings if f.get("file_name")))
        if all_files:
            reply = "# All File Names\n" + "\n".join(f"- {f}" for f in all_files)
        else:
            reply = "_No files found in the current evidence._"
        return {"reply": reply}

    # 4. Full findings JSON
    if re.search(r'(list|show|all) (records|rows|details|findings)', user_query):
        pretty = json.dumps(findings, indent=2, ensure_ascii=False)
        reply = f"```json\n{pretty}\n```"
        return {"reply": reply}

    # 5. Chatbot LLM (fallback to narrative/generation for other queries)
    persona_instructions = {
        "Risk Auditor": (
            "You answer as a highly professional compliance/risk auditor. Use detailed, well-structured markdown with headings, bullets, bold text, and tables if appropriate."
        ),
        "Cybersecurity": (
            "You are an experienced cyber defense expert. ALWAYS use markdown structure in your response (use code blocks, tables, strong callouts where useful)."
            " IMPORTANT: If the 'Available File Evidence' section below says there is NO EVIDENCE or NO MATCHING results, you must answer honestly: 'No matching file changes found for your request.' But include content in your response that is not reliant on data evidence."
            " DO NOT invent evidence data; summarize in the context of request."
        ),
        "Board Member": (
            "You are a board-level advisor. Provide summaries, lists, and strong section headings in Markdown for business readability."
            " IMPORTANT: If the 'Available File Evidence' section below says there is NO EVIDENCE or NO MATCHING results, you must answer honestly: 'No matching file changes found for your request.' But include content in your response that is not reliant on data evidence."
            " DO NOT invent evidence data; summarize in the context of request."
        ),
        "All": (
            "You are CybersecAI, combining compliance, cyber, and business expertise. Provide an integrated response from both compliance and security perspectives, always using clean, structured Markdown."
            " IMPORTANT: If the 'Available File Evidence' section below says there is NO EVIDENCE or NO MATCHING results, you must answer honestly: 'No matching file changes found for your request.' But include content in your response that is not reliant on data evidence."
            " DO NOT invent evidence data; summarize in the context of request."
        ),
        "Auditor/Security": (
            "You are a combined compliance auditor and cybersecurity analyst. Surface forensic traces and risk maps, using structured Markdown for clarity."
            " IMPORTANT: If the 'Available File Evidence' section below says there is NO EVIDENCE or NO MATCHING results, you must answer honestly: 'No matching file changes found for your request.' But include content in your response that is not reliant on data evidence."
            " DO NOT invent evidence data; summarize in the context of request."
        ),
        "default": (
            "You are CybersecAI, an AI chat advisor. Always return clean, structured Markdown using headers, lists, bold, etc."
            " IMPORTANT: If the 'Available File Evidence' section below says there is NO EVIDENCE or NO MATCHING results, you must answer honestly: 'No matching file changes found for your request.' But include content in your response that is not reliant on data evidence."
            " DO NOT invent evidence data; summarize in the context of request."
        )
    }
    system_content = persona_instructions.get(req.persona, persona_instructions["default"])

    date_bounds = parse_date_from_query(req.query)
    filtered_evidence = []
    if date_bounds:
        for entry in findings:
            lastmod = entry.get('last_modified')
            if lastmod:
                try:
                    dt = dateutil.parser.parse(lastmod)
                    if date_bounds[0] <= dt <= date_bounds[1]:
                        filtered_evidence.append(entry)
                except Exception:
                    continue
    elif re.search(r'(last|recent)\s+\d+\s+(file|change)', req.query, re.I):
        filtered_evidence = findings[-100:]
    else:
        filtered_evidence = findings[-100:]

    # Replace evidence_context to show the storage-backend summary!
    evidence_context = storage_md

    #Optionally: Add additional listing for the specific files in the timeframe, if desired:
    if filtered_evidence:
        evidence_context += "\n#### Selected Files in Range:\n"
        for f in filtered_evidence[:300]:
            try:
                fname = f.get('file_name') or f.get('filename') or "[unknown file]"
                lastmod = f.get('last_modified') or f.get('lastModified') or ""
                risk = f.get("overall_risk_rating") or ""
                evidence_context += f"- {fname} ({lastmod}) Risk: {risk}\n"
            except Exception:
                continue

    if not filtered_evidence:
        evidence_context += "\n[NO MATCHING HIGH-RISK FILE CHANGES for your account for the specified date/range. No results found. AI MUST NOT INVENT FILES.]"

    # ---- Assemble chat history and pass evidence_context to LLM ----
    full_history = req.messages[-2:]
    conversation = [
        {
            "role": "system",
            "content": (
                system_content +
                "\n\nRespond ALWAYS in Markdown and NEVER in plain text or code - use headers, lists, bold, call-outs. Answer queries using the full evidence below:"
            )
        }
    ]
    for m in full_history:
        if m.get("role") and m.get("content"):
            conversation.append({"role": m["role"] if m["role"] in ["user", "assistant"] else "user", "content": m["content"]})
    conversation.append({"role": "user", "content": f"Available File Evidence:\n{evidence_context}\n\nUser Query: {req.query}"})

    try:
        resp = client.chat.completions.create(
            model="gpt-4.1",
            messages=conversation,
            temperature=0.2,
        )
        text = resp.choices[0].message.content
        output_guardrails(text)
        return {"reply": text}
    except Exception as e:
        return {"reply": f"AI error: {str(e)}"}

# === PROVIDE PROMPTS AND LOGIC for All Other Agents, Unmodified (compliance, audit, etc.)
def agent_compliance_advisor(data: Dict[str, Any]):
    prompt = f"""
You are the cybersecai.io compliance expert AI Agent. Your reply must follow the professional standards and compliance guardrails defined by cybersecai.io where applicable.

Inputs:
Standard: {data.get('standard')}
Jurisdiction: {data.get('jurisdiction')}
Details: {data.get('requirement_notes', '')}
Event: {data.get('event_type', '')}
Data: {json.dumps(data.get('data', {}), ensure_ascii=False)}

Instructions:
- Your risk assessment and recommendations are provided as the expert opinion of the cybersecai.io compliance AI Agent, in line with cybersecai.io's guardrails and with full impartiality.
- Clearly indicate this at the end of your response as a formal disclaimer.

Please:
1: Score the privacy risk (LOW, MEDIUM, HIGH).
2: Recommend next action(s) (internal_report, notify_authority, communicate_subjects, public_communication, etc).
3: Identify and summarize any other legal or regulatory compliance obligations that may apply in {data.get('jurisdiction')}, beyond {data.get('standard')} (including sector-specific or cross-jurisdictional duties).
4: Generate a notification/report letter that satisfies both the requirements for this standard and broader obligations in the jurisdiction.
5: Write your full response and letter in polished, formal English and use Markdown formatting. Include risk, action, decision_reason, notification_letter (if required), **and a final section clearly stating:**  
"_This opinion is produced by the cybersecai.io compliance expert AI Agent and is based on its programmed guardrails and up-to-date regulatory analysis._"
    """.strip()
    try:
        resp = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are an automated compliance advisor."},
                {"role": "user", "content": prompt}
            ]
        )
        text = resp.choices[0].message.content
        output_guardrails(text)
        return {"markdown": text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Compliance LLM failed: {str(e)}")

def agent_audit(data: Dict[str, Any]):
  
    prompt = (
        f"You are a highly professional internal auditor preparing a risk report for the Board of Directors in {data.get('region')}. "
        "You are provided with AI-generated file risk summaries from the cybersecai.io platform, each record assessed as HIGH overall risk. "
        "Each file record contains information such as file name, the 'auditor_proposed_action', last modified time, data classification, and business area."
        "You must analyze, compare, and highlight high-risk file changes and trends across time windows: "
        "specifically, (1) the last 1 month, and (2) the prior 2-month window (excluding the most recent month). "
        "Help the Board understand what risks are *new*, what has *persisted*, and what is *escalating or recurring*."
        "\n\n"
        f"Raw JSON data to audit is below:\n\n{data.get('json_data')}\n\n"
        "Your report MUST use Markdown headings, clear Board-level language, and be structured as follows:\n"
        "1. ### Executive Summary\n"
        "   - Concisely describe the organization's overall exposure, and highlight any sudden changes or escalations in the last month.\n"
        "2. ### Trends and Red Flags: Last 1 Month\n"
        "   - Analyze and bullet key risk trends in just the last month (by file type, department, business area).\n"
        "   - Markdown table: list files changed/flagged in the last month, with columns for file name, last modified, data type, proposed action, and any Board-urgent flags (use bold/callouts if needed).\n"
        "   - List any *new* or *escalated* risks not seen in the previous months.\n"
        "3. ### Trends and Red Flags: 1-2 Months Ago\n"
        "   - Analyze and bullet patterns in the window 1-2 months ago. Show its own Markdown table (same columns).\n"
        "   - Summarize risks that have persisted (appearing in both tables), and risks that were resolved or remediated.\n"
        "4. ### Comparative Analysis\n"
        "   - Compare the two time periods. Clearly highlight:\n"
        "     * New/recurring risks in the last month\n"
        "     * Which risks have emerged, worsened, or faded\n"
        "     * Any areas where risk control has improved or slipped\n"
        "     * Any files/actions needing Board-level escalation or urgent review\n"
        "5. ### Governance & Compliance Implications\n"
        "   - Describe, for {data.get('region')}, what legal/reporting obligations arise out of the last 1 month’s findings (especially if new/escalated). Flag if formal authority notification or public statement may be required.\n"
        "6. ### Board-Ready Recommendations\n"
        "   - Numbered, actionable next steps for the Board, using Board-resolution language (who, what, urgent deadlines)."
        "7. ### Risk Evolution Tables\n"
        "   - Render side-by-side (or one after the other) Markdown tables summarizing high-risk file counts by key action/type for each period (last month vs prior).\n"
        "\nStrictly do NOT repeat or dump raw JSON—synthesize, cluster, and write for the Board’s review. Use Markdown headings (### ...), bullets, and tables throughout. Highlight urgent items in bold or with callouts!"
    )
    try:
        resp = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are a world class internal auditor."},
                {"role": "user", "content": prompt}
            ]
        )
        output = resp.choices[0].message.content
        output_guardrails(output)
        return {"markdown": output}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audit LLM failed: {str(e)}")

def dict_to_markdown_table(d, title=""):
    # d: dict mapping strings to numbers (typically)
    lines = []
    if title:
        lines.append(f"### {title}")
    # Table header
    lines.append("| Item | Count |")
    lines.append("|------|-------|")
    for k, v in d.items():
        lines.append(f"| {k} | {v} |")
    return "\n".join(lines)

def list_to_markdown(lst, title=""):
    # lst: list of values/strings
    lines = []
    if title:
        lines.append(f"### {title}")
    for item in lst:
        lines.append(f"- {item}")
    return "\n".join(lines)


def agent_findings_facts(data: Dict[str, Any]):
    """
    Returns readable summaries for:
      - 'How many high risk files overall?'
      - 'How many high risk files by jurisdiction?'
      - 'Show me all files found in <backend>'
    Otherwise, returns pretty JSON.
    """
    op = data.get("operation", "")
    args = data.get("args", {}) if data.get("args") is not None else {}
    config_ids = data.get("config_ids")
    #findings = gather_high_risk_findings(config_ids)
    user_query = (data.get("query") or op or "").lower()

    # Use all findings for medium/low/all risk, else only high!
    if any(kw in user_query for kw in ["medium risk", "low risk", "all risk", "any risk"]):
        findings = gather_all_findings(config_ids)
    else:
        findings = gather_high_risk_findings(config_ids)

        # --- Data classification count summary ---
    if re.search(r"(how many|count|number of).*(file|files).*(data classification|classification)", user_query):
        class_count = {}
        for f in findings:
            cls = (f.get("data_classification") or "").strip()
            if cls:
                class_count[cls] = class_count.get(cls, 0) + 1
            else:
                class_count["[Unclassified]"] = class_count.get("[Unclassified]", 0) + 1
        if not class_count:
            reply = "_No data classifications found in current high risk files._"
        else:
            reply = "### High Risk File Count by Data Classification\n\n"
            reply += "| Data Classification | File Count |\n"
            reply += "|---------------------|------------|\n"
            for k, v in class_count.items():
                reply += f"| {k} | {v} |\n"
        return {"reply": reply, "raw": class_count, "operation": op, "args": args}


    def md_table(headers, rows, title=""):
        out = []
        if title: out.append(f"### {title}")
        out.append("| " + " | ".join(headers) + " |")
        out.append("|" + "|".join(["-"*len(h) for h in headers]) + "|")
        for row in rows:
            out.append("| " + " | ".join(str(x) for x in row) + " |")
        return "\n".join(out)


    # --- Data classification count summary ---
    if re.search(r"(how many|count|number of).*(file|files).*(data classification|classification)", user_query):
        class_count = {}
        for f in findings:
            cls = (f.get("data_classification") or "").strip()
            if cls:
                class_count[cls] = class_count.get(cls, 0) + 1
            else:
                class_count["[Unclassified]"] = class_count.get("[Unclassified]", 0) + 1
        if not class_count:
            reply = "_No data classifications found in current high risk files._"
        else:
            reply = "### High Risk File Count by Data Classification\n\n"
            reply += "| Data Classification | File Count |\n"
            reply += "|---------------------|------------|\n"
            for k, v in class_count.items():
                reply += f"| {k} | {v} |\n"
        return {"reply": reply, "raw": class_count, "operation": op, "args": args}

    # How many files are rated as high risk overall?
    if ("how many" in user_query or "count" in user_query) and ("high risk" in user_query) and (
         "overall" in user_query or "risk rating" in user_query or "overall risk" in user_query):
        n = len(findings)
        reply = md_table(
           headers=["Overall Risk Rating", "Number of Files"],
           rows=[["High", n]],
           title="High Risk File Summary"
        )
        #return {"reply": reply, "raw": n, "operation": op, "args": args}
        return reply

    # How many high risk files by jurisdiction?
    if ("how many" in user_query or "count" in user_query) and (
        "jurisdiction" in user_query or "each jurisdiction" in user_query or "by jurisdiction" in user_query
    ):
        # Count files per jurisdiction - each file can count for more than one
        jurisdiction_counts = {}
        for file in findings:
            if isinstance(file.get("compliance_findings"), list):
                jurisdictions = set()
                for cf in file["compliance_findings"]:
                    j = cf.get("jurisdiction")
                    if j:
                        jurisdictions.add(j)
                for j in jurisdictions:
                    jurisdiction_counts[j] = jurisdiction_counts.get(j, 0) + 1
        if not jurisdiction_counts:
            reply = "_No jurisdictions detected in current high-risk findings._"
        else:
            rows = [(j, c) for j, c in jurisdiction_counts.items()]
            reply = md_table(
                headers=["Jurisdiction", "High Risk File Count"],
                rows=rows,
                title="High Risk Files by Jurisdiction"
            )
        return {"reply": reply, "raw": jurisdiction_counts, "operation": op, "args": args}

    # Show all files found in <backend> (SMB, S3, NFS, M365, etc.)
    backend_names = ["smb", "s3", "nfs", "m365"]
    for b in backend_names:
        if f"all files" in user_query and b in user_query:
            # Make a flat table of files with data_source == b (case-insensitive)
            matching = [f for f in findings if (f.get("data_source","").lower() == b)]
            if not matching:
                reply = f"_No high risk files found in the {b.upper()} backend._"
            else:
                headers = ["File Name", "Last Modified", "Data Classification", "Overall Risk", "File Path"]
                rows = [
                    [
                        f.get("file_name", ""),
                        f.get("last_modified", ""),
                        f.get("data_classification", ""),
                        f.get("overall_risk_rating", ""),
                        f.get("full_path", f.get("file_path", "")),
                    ]
                    for f in matching
                ]
                reply = md_table(headers, rows, f"High Risk Files in {b.upper()} Backend")
            return {"reply": reply, "raw": matching, "operation": op, "args": args}

    # List all details for every file currently rated high risk
    # if re.search(r'(list|show|all).*(detail|record|full|everything)(.*high risk)?', user_query):
    #     markdown = format_high_risk_files_markdown(findings)
    #     return {"reply": markdown, "raw": findings, "operation": op, "args": args}

    if re.search(r'(list|show|all).*(detail|record|full|everything).*(high risk)', user_query, re.I):
        markdown = format_high_risk_files_markdown(findings)
        return {"reply": markdown, "raw": findings, "operation": op, "args": args}

    elif re.search(r'(list|show|all).*(detail|record|full|everything).*(medium risk)', user_query, re.I):
        markdown = format_medium_risk_files_markdown(findings)
        return {"reply": markdown, "raw": findings, "operation": op, "args": args}

    elif re.search(r'(list|show|all).*(detail|record|full|everything).*(risk)', user_query, re.I):
        markdown = format_all_risks_files_markdown(findings)
        return {"reply": markdown, "raw": findings, "operation": op, "args": args}


    # You can add further custom clauses here if needed!

    # Default: fallback to pretty JSON
    pretty = json.dumps(findings, indent=2, ensure_ascii=False)
    reply = f"```json\n{pretty}\n```"
    return {"reply": reply, "raw": findings, "operation": op, "args": args}



def agent_policy_enforce(data: Dict[str, Any]):
    # From your previous code, nothing omitted
    changed_files, policy_actions, siem_events = [], [], []
    files = data['files']
    policies = data['policies']
    siem_url = data.get('siem_url', "")

    for fname, info in files.items():
        if info.get('changed'):
            changed_files.append(fname)
            compliance_decision = None
            if info.get('risk', '').upper() == "HIGH":
                comp_req = {
                    "standard": policies.get('standard', 'GDPR'),
                    "jurisdiction": policies.get('jurisdiction', 'Europe'),
                    "requirement_notes": policies.get('notes', ''),
                    "event_type": "Sensitive File Change",
                    "data": info
                }
                comp_res = agent_compliance_advisor(comp_req)
                policy_actions.append({
                    "file": fname, "compliance_decision": comp_res
                })
                compliance_decision = comp_res
            if info.get('policy_required'):
                policy_actions.append({
                    "file": fname, "action_taken": f"Policy {policies.get('enforce_type', 'Lock')} applied"
                })
            if siem_url:
                siem_event = {
                    "to": siem_url,
                    "event": {
                        "file": fname,
                        "delta": info.get('delta'),
                        "compliance_decision": compliance_decision
                    }
                }
                siem_events.append(siem_event)
    return {
        "changed_files": changed_files,
        "policy_actions": policy_actions,
        "siem_events": siem_events
    }

# ---- Agent Registry & "What fields are required" ----
AGENT_SCHEMAS = {
    "compliance": {
        "required": ["standard", "jurisdiction", "event_type", "data"],
    },
    "audit": {
        "required": ["region", "json_data"],
    },
    "policy_enforce": {
        "required": ["files", "policies"],
    },
    "findings_facts": {
        "required": ["operation", "config_ids", "args"],  
    },
    "chatbot": {
        "required": ["persona", "query", "messages", "config_ids"],
    },
}

AGENT_HANDLERS = {
    'compliance': agent_compliance_advisor,
    'audit': agent_audit,
    'policy_enforce': agent_policy_enforce,
    'findings_facts': agent_findings_facts,
    'chatbot': agent_chatbot,
}





def substrings_in(text, keywords):
    return any(kw in text for kw in keywords)

def infer_agent_and_missing(data: Dict[str, Any], user_query: str):
    q = user_query.lower()
    # Robust regex-based detection
    

    compliance_regex = re.compile(
        r'(compliance|gdpr|privacy act|ccpa|cpra|australian privacy|'
        r'personal information|regulation|breach notification|'
        r'sensitive file|privacy violation|data breach|notifiable breach|'
        r'loss of data|accident|medical info|pii|'
        r'employee data|data leak|sensitive data|exposed pii)'
    )
    audit_regex = re.compile(
        r'(audit|board review|board report|summary for board|executive summary|board-level|material risk|governance|internal auditor|for board)'
    )
    policy_regex = re.compile(
        r'(policy|enforce|remediation|lock this file|quarantine|auto-encrypt|'
        r'access control|auto-classification|auto-remediation|'
        r'zero trust|secure by design)'
    )

    findings_facts_regex = re.compile(
        r'(how many|count|number of|list|show me|top \d+|latest \d+|summary of|breakdown).*\b(file|risk|m365|s3|smb|classification|pii|phi|high risk)\b',
        flags=re.I)

    # KEY: persist agent through context until all its fields are filled!
    agent = None
    # If agent already in context, keep using it
    if 'agent' in data and data['agent'] in AGENT_SCHEMAS:
        agent = data['agent']
    else:
        # Otherwise, use most recent user query text for selection
        if audit_regex.search(q):
            agent = 'audit'
        elif compliance_regex.search(q):
            agent = 'compliance'
        elif policy_regex.search(q):
            agent = 'policy_enforce'
        elif findings_facts_regex.search(q):
            agent = 'findings_facts'
        else:
            agent = 'chatbot'

    # Save agent to context for multi-turn lock-in
    data['agent'] = agent

    # No user prompt ever for json_data-- auto slotfill in orchestrate!
    required = AGENT_SCHEMAS[agent]['required']
    missing = [f for f in required if f not in data]
    #missing = [f for f in required if f not in data or not data[f]]

    # For audit agent, remove 'json_data' from missing if gather_high_risk_findings will fill it
    if agent == "audit" and "json_data" in missing:
        config_ids = data.get("config_ids")
        if config_ids is not None and len(config_ids) > 0:
            missing = [f for f in missing if f != "json_data"]

    return agent, missing

def get_session_state(session_id:str) -> dict:
    return session_mind.setdefault(session_id, {})

def update_session_state(session_id:str, **kwargs):
    state = get_session_state(session_id)
    state.update(kwargs)
    session_mind[session_id] = state
    return state

# --- API: Autonomous Orchestrator ---
class UserChatReq(BaseModel):
    user_query: str
    prior_context: Dict[str, Any] = {}
    session_id: Optional[str] = None

def markdown_to_docx(md_content, docx_path):
    import re
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    lines = md_content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Markdown Heading
        if re.match(r"^#{1,6} ", line):
            level = len(line.split(" ")[0])
            doc.add_heading(line[level+1:].strip(), min(level-1, 4))

        # Markdown Table
        elif line.strip().startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            headers = [c.strip() for c in table_lines[0].strip().strip('|').split('|')]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            for idx, text in enumerate(headers):
                table.rows[0].cells[idx].text = text
            for trow in table_lines[2:]:
                vals = [c.strip() for c in trow.strip().strip('|').split('|')]
                row = table.add_row()
                for idx, cell in enumerate(vals):
                    row.cells[idx].text = cell

        # Bullet Point (with bold detection)
        elif re.match(r"^\s*[-*+] ", line) or re.match(r"^• ", line):
            items = []
            while i < len(lines) and (re.match(r"^\s*[-*+] ", lines[i]) or re.match(r"^• ", lines[i])):
                bullet = re.sub(r"^(\s*[-*+]|•)\s*", '', lines[i]).strip()
                items.append(bullet)
                i += 1
            i -= 1
            idx = 0
            while idx < len(items):
                bold_heading = re.match(r"^\*\*(.+?)\*\*:?$", items[idx])
                if bold_heading:
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(bold_heading.group(1) + (":" if items[idx].endswith(":") else ""))
                    run.bold = True
                    idx += 1
                    subidx = idx
                    while subidx < len(items) and not re.match(r"^\*\*.+?\*\*", items[subidx]):
                        para2 = doc.add_paragraph(items[subidx], style='List Bullet 2')
                        subidx += 1
                    idx = subidx
                else:
                    doc.add_paragraph(items[idx], style='List Bullet')
                    idx += 1

        # Numbered list
        elif re.match(r"^\s*\d+\.", line):
            while i < len(lines) and re.match(r"^\s*\d+\.", lines[i]):
                txt = re.sub(r"^\s*\d+\.\s*", '', lines[i]).strip()
                doc.add_paragraph(txt, style='List Number')
                i += 1
            i -= 1

        # Horizontal rule
        elif re.match("^(-{3,}|_{3,}|\\*{3,})$", line.strip()):
            doc.add_paragraph().add_run().add_break()

        # Blockquote
        elif line.strip().startswith(">"):
            quote = line.lstrip("> ").strip()
            para = doc.add_paragraph()
            run = para.add_run(quote)
            run.italic = True
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Regular paragraph
        elif line.strip() != "":
            para = doc.add_paragraph(line)

        i += 1
    doc.save(docx_path)



def get_csv_download_link(findings):
    import csv, os
    base_dir = "/home/cybersecai/htdocs/www.cybersecai.io/webhook/agenticai/tmp_csv"
    os.makedirs(base_dir, exist_ok=True)  # Always create the tmp_csv folder if it doesn't exist

    csv_name = f"highrisk_{uuid.uuid4().hex}.csv"
    csv_path = os.path.join(base_dir, csv_name)
    if findings:
        fields = set()
        for f in findings:
            fields.update(f.keys())
        csv_rows = []
        for f in findings:
            row = dict(f)
            if isinstance(row.get('compliance_findings'), list):
                row['compliance_findings'] = '; '.join(str(x) for x in row['compliance_findings'])
            if isinstance(row.get('permissions'), list):
                row['permissions'] = '; '.join(str(x) for x in row['permissions'])
            csv_rows.append(row)
        with open(csv_path, "w", newline="", encoding="utf-8") as fcsv:
            writer = csv.DictWriter(fcsv, fieldnames=sorted(fields))
            writer.writeheader()
            writer.writerows(csv_rows)
    csv_url = f"/download_csv?file={csv_name}"
    return csv_url


def robust_write_docx_file(md_text: str) -> str:
    import traceback
    base_dir = "/home/cybersecai/htdocs/www.cybersecai.io/webhook/agenticai/tmp_csv"
    os.makedirs(base_dir, exist_ok=True)
    docx_name = f"chat_output_{uuid.uuid4().hex}.docx"
    docx_path = os.path.join(base_dir, docx_name)
    try:
        markdown_to_docx(md_text, docx_path)
        return f"/download_docx?file={docx_name}"
    except Exception as ex:
        log_to_laravel(f"[markdown_to_docx error] {repr(ex)}\n{traceback.format_exc()}")
        return ""

@app.get("/download_docx")
def download_docx(file: str):
    base_dir = "/home/cybersecai/htdocs/www.cybersecai.io/webhook/agenticai/tmp_csv"
    file_path = os.path.join(base_dir, file)
    if not os.path.exists(file_path) or not file.endswith('.docx'):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        file_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=os.path.basename(file_path)
    )

@app.get("/download_csv")
def download_csv(file: str):
    base_dir = "/home/cybersecai/htdocs/www.cybersecai.io/webhook/agenticai/tmp_csv"
    file_path = os.path.join(base_dir, file)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, media_type='text/csv', filename=os.path.basename(file_path))

@app.post("/agentic/auto_orchestrate")
def auto_orchestrate(req: UserChatReq):
    session_id = req.session_id or str(uuid4())
    state = get_session_state(session_id)
    context = {**(state.get('context', {})), **req.prior_context, "query": req.user_query}

    agent, missing = infer_agent_and_missing(context, req.user_query)

    # Automatically fill placeholder args if agent is findings_facts and args is missing
    if agent == "findings_facts" and "args" not in context:
        context["args"] = {}

    # For audit agent: slot-fill json_data if needed
    if agent == "audit":
        if not context.get("json_data"):
            config_ids = context.get("config_ids")
            if config_ids:
                findings = gather_high_risk_findings(config_ids)
                context["json_data"] = json.dumps(findings[:50])  # Limit size
        if "json_data" in missing and context.get("json_data"):
            missing = [f for f in missing if f != "json_data"]

    # Always ensure messages is array
    if not context.get("messages"):
        context["messages"] = [{"role": "user", "content": context["query"]}]
    elif isinstance(context.get("messages"), str):
        context["messages"] = [{"role": "user", "content": context["messages"]}]
    elif isinstance(context.get("messages"), list) and len(context["messages"]) == 0:
        context["messages"] = [{"role": "user", "content": context["query"]}]

    log_to_laravel(f"AutoOrchestrator: intent={agent}, missing={missing}, context={json.dumps(context)[:400]}")

    required = AGENT_SCHEMAS[agent]['required']
    # Final check: Only check for PRESENCE of keys, not truthiness
    missing = [f for f in required if f not in context]

    if agent == "findings_facts" and "args" in missing:
        # Safety: forcibly add empty args if still missing (shouldn't happen!)
        context['args'] = {}
        missing = [f for f in missing if f != 'args']

    if agent == "audit" and "json_data" in missing and context.get("json_data"):
        missing = [f for f in missing if f != "json_data"]

    if missing:
        question = f"To process your request as a {agent.replace('_',' ')}, I need the following: " + ", ".join(missing)
        update_session_state(session_id, context=context, agent=agent, missing=missing)
        return {
            "pending": True,
            "question": question,
            "session_id": session_id
        }
    else:
        handler = AGENT_HANDLERS[agent]
        try:
            result = handler(context)
            log_to_laravel(f"AutoOrchestrator Result {agent}: {str(result)[:450]}")
        except Exception as ex:
            result = f"Error running agent: {ex}"
            log_to_laravel(str(result))
        update_session_state(session_id, context=context, agent=agent, result=result)
    

    # ---- Universal DOCX and CSV handling for download links in all responses ----
    config_ids = context.get("config_ids")
    csv_url = ""
    docx_url = ""
    docx_source = ""

    if config_ids:

        if agent in ("findings_facts", "chatbot"):
            user_query = context.get("query") or ""
            if any(kw in user_query.lower() for kw in ["medium risk", "low risk", "all risk", "any risk"]):
                findings = gather_all_findings(config_ids)
            else:
                findings = gather_high_risk_findings(config_ids)
        else:
            findings = gather_high_risk_findings(config_ids)
        csv_url = get_csv_download_link(findings)

        # Only generate docx for agents or result types needing a formal document
        agent_for_docx = agent in ("audit", "compliance", "chatbot")
        # Only proceed if the output seems markdown or main report style (not a pure number/array)
        if agent_for_docx:
            if isinstance(result, dict):
                docx_source = result.get("reply") or result.get("markdown") or ""
            elif isinstance(result, str):
                docx_source = result
            # Heuristic: skip for numerical or pure list result
            if docx_source and docx_source.strip() and len(docx_source) > 30 and not docx_source.strip().isdigit():
                try:
                    docx_url = robust_write_docx_file(docx_source)
                except Exception as ex:
                    log_to_laravel(f"Docx generation skipped or failed for non-markdown result: {repr(ex)}")

    download_docx_md = f"\n\n---\n[⬇️ Download this report as Word (docx)]({docx_url})" if docx_url else ""
    download_csv_md = f"\n\n---\n[⬇️ Download all high risk findings as CSV]({csv_url})" if csv_url else ""

    if isinstance(result, dict):
        if "reply" in result and isinstance(result["reply"], str):
            result["reply"] += download_docx_md + download_csv_md
        if "markdown" in result and isinstance(result["markdown"], str):
            result["markdown"] += download_docx_md + download_csv_md
    elif isinstance(result, str):
        result += download_docx_md + download_csv_md
    # ---- END Universal Download Links block ----

    return {
        "pending": False,
        "result": result,
        "session_id": session_id
    }




if __name__ == "__main__":
    import uvicorn
    uvicorn.run("agentic_orchestrator_service:app", host="127.0.0.1", port=8223)